"""
This sample, non-production-ready application will produce Word Document transcriptions using automatic speech
recognition from Amazon Transcribe, and handles all processing modes of Amazon Transcribe in terms of diarization:
speaker-separated audio, channel-separated audio, or Call Analytics audio.  The application requires the following
non-standard python libraries to be installed:

- python-docx
- scipy
- matplotlib

© 2021 Amazon Web Services, Inc. or its affiliates. All Rights Reserved.
This AWS Content is provided subject to the terms of the AWS Customer Agreement available at
http://aws.amazon.com/agreement or other written agreement between Customer and either
Amazon Web Services, Inc. or Amazon Web Services EMEA SARL or both.
"""

from docx import Document
from docx.shared import Cm, Mm, Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_SECTION
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from pathlib import Path
from time import perf_counter
from scipy.interpolate import make_interp_spline
import urllib.request
import json
import datetime
import matplotlib.pyplot as plt
import matplotlib.ticker as ticker
import numpy as np
import statistics
import os
import boto3
import argparse
from io import BytesIO


# Common formats and styles
CUSTOM_STYLE_HEADER = "CustomHeader"
TABLE_STYLE_STANDARD = "Light List"
CATEGORY_TRANSCRIPT_BG_COLOUR = "EEFFFF"
CATEGORY_TRANSCRIPT_FG_COLOUR = RGBColor(0, 128, 255)
ALTERNATE_ROW_COLOUR = "F0F0F0"
BAR_CHART_WIDTH = 1.0

# Column offsets in Transcribe output document table
COL_STARTTIME = 0
COL_ENDTIME = 1
COL_SPEAKER = 2
COL_SENTIMENT = 3
COL_CONTENT = 4

# Comprehend Sentiment helpers - note, if a language code in Comprehend has multiple suffixed versions
# then the suffixed versions MUST be defined in the language list BEFORE the base one; e.h. "zh-TW" before "zh"
MIN_SENTIMENT_LENGTH = 16
MIN_SENTIMENT_NEGATIVE = 0.4
MIN_SENTIMENT_POSITIVE = 0.6
SENTIMENT_LANGUAGES = ["en", "es", "fr", "de", "it", "pt", "ar", "hi", "ja", "ko", "zh-TW", "zh"]

# Image download URLS
IMAGE_URL_BANNER = "https://raw.githubusercontent.com/aws-samples/amazon-transcribe-output-word-document/main/images/banner.png"
IMAGE_URL_SMILE = "https://raw.githubusercontent.com/aws-samples/amazon-transcribe-output-word-document/main/images/smile.png"
IMAGE_URL_FROWN = "https://raw.githubusercontent.com/aws-samples/amazon-transcribe-output-word-document/main/images/frown.png"
IMAGE_URL_NEUTRAL = "https://raw.githubusercontent.com/aws-samples/amazon-transcribe-output-word-document/main/images/neutral.png"

# Definitions to use whilst scanning summarisation data
CALL_SUMMARY_MAP = [
    {"Field": "segmentIssuesDetected", "Title": "Issues Detected", "Color": "FF3333"},
    {"Field": "segmentActionItemsDetected", "Title": "Action Items Detected", "Color": "FFB266"},
    {"Field": "segmentOutcomesDetected", "Title": "Outcomes Detected", "Color": "66CC00"}
]

# Additional Constants
START_NEW_SEGMENT_DELAY = 2.0       # After n seconds pause by one speaker, put next speech in new segment


class SpeechSegment:
    """ Class to hold information about a single speech segment """
    def __init__(self):
        self.segmentStartTime = 0.0
        self.segmentEndTime = 0.0
        self.segmentSpeaker = ""
        self.segmentText = ""
        self.segmentConfidence = []
        self.segmentSentimentScore = -1.0    # -1.0 => no sentiment calculated
        self.segmentPositive = 0.0
        self.segmentNegative = 0.0
        self.segmentIsPositive = False
        self.segmentIsNegative = False
        self.segmentAllSentiments = []
        self.segmentLoudnessScores = []
        self.segmentInterruption = False
        self.segmentIssuesDetected = []
        self.segmentActionItemsDetected = []
        self.segmentOutcomesDetected = []


def convert_timestamp(time_in_seconds):
    """
    Function to help convert timestamps from s to H:M:S:MM

    :param time_in_seconds: Time in seconds to be displayed
    :return: Formatted string for this timestamp value
    """
    timeDelta = datetime.timedelta(seconds=float(time_in_seconds))
    tsFront = timeDelta - datetime.timedelta(microseconds=timeDelta.microseconds)
    tsSmall = timeDelta.microseconds
    return str(tsFront) + "." + str(int(tsSmall / 10000))


def get_text_colour_analytics_sentiment(score):
    """
    Returns RGB code text to represent the strength of negative or positive sentiment

    :param score: Sentiment score in range +/- 5.0
    :return: Background RGB colour text string to use in sentiment text
    """
    # Get our score into the range [0..4], which is our shade 'strength' - higher => brighter shade
    truncated = min(abs(int(score)), 4)
    col_shade = (4 - truncated) * 51

    if score >= 0:
        # Positive sentiment => Green shade
        background_colour = "{0:0>2X}{1:0>2X}{2:0>2X}".format(col_shade, 255, col_shade)
    else:
        # Negative sentiment => Red shade
        background_colour = "{0:0>2X}{1:0>2X}{2:0>2X}".format(255, col_shade, col_shade)

    return background_colour


def set_table_row_bold(row, bold):
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = bold

def set_transcript_text_style(run, force_highlight, confidence=0.0, rgb_color=None):
    """
    Sets the colour and potentially the style of a given run of text in a transcript.  You can either
    supply the hex-code, or base it upon the confidence score in the transcript.

    :param run: DOCX paragraph run to be modified
    :param force_highlight: Indicates that we're going to forcibly set the background colour
    :param confidence: Confidence score for this word, used to dynamically set the colour
    :param rgb_color: Specific colour for the text
    """

    # If we have an RGB colour then use it
    if rgb_color is not None:
        run.font.color.rgb = rgb_color
    else:
        # Set the colour based upon the supplied confidence score
        if confidence >= 0.90:
            run.font.color.rgb = RGBColor(0, 0, 0)
        elif confidence >= 0.5:
            run.font.color.rgb = RGBColor(102, 51, 0)
        else:
            run.font.color.rgb = RGBColor(255, 0, 0)

    # Apply any other styles wanted
    if confidence == 0.0:
        # Call out any total disasters in bold
        run.font.bold = True

    # Force the background colour if required
    if force_highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def write_transcribe_text(output_table, sentiment_enabled, analytics_mode, speech_segments, keyed_categories):
    """
    Writes out each line of the transcript in the Word table structure, optionally including sentiments

    :param output_table: Word document structure to write the table into
    :param sentiment_enabled: Flag to indicate we need to show some sentiment
    :param analytics_mode: Flag to indicate we're in Analytics mode, not Standard
    :param speech_segments: Turn-by-turn speech list
    :param keyed_categories: List of categories identified at any timestamps
    """

    # Load our image files if we have sentiment enabled
    if sentiment_enabled:
        png_smile = load_image(IMAGE_URL_SMILE)
        png_frown = load_image(IMAGE_URL_FROWN)
        png_neutral = load_image(IMAGE_URL_NEUTRAL)
        content_col_offset = 0
    else:
        # Ensure we offset the CONTENT column correctly due to no sentiment
        content_col_offset = -1

    # Create a row populate it for each segment that we have
    shading_reqd = False
    for segment in speech_segments:
        # Before we start, does an angory start at this time?
        start_in_millis = segment.segmentStartTime * 1000.0
        end_in_millis = segment.segmentEndTime * 1000.0
        if start_in_millis in keyed_categories:
            insert_category_row(content_col_offset, keyed_categories, output_table, start_in_millis)
            keyed_categories.pop(start_in_millis)

        # Start with the easy stuff
        row_cells = output_table.add_row().cells
        row_cells[COL_STARTTIME].text = convert_timestamp(segment.segmentStartTime)
        row_cells[COL_ENDTIME].text = f"{(segment.segmentEndTime - segment.segmentStartTime):.1f}s"
        row_cells[COL_SPEAKER].text = segment.segmentSpeaker

        # Mark the start of the turn as INTERRUPTED if that's the case
        if segment.segmentInterruption:
            run = row_cells[COL_CONTENT + content_col_offset].paragraphs[0].add_run("[INTERRUPTION]")
            set_transcript_text_style(run, True, confidence=0.0)
            row_cells[COL_CONTENT + content_col_offset].paragraphs[0].add_run(" ")

        # Summarised data blocks are in order - pick out the first for each of our
        # types, as well as getting list of the remaining ones for this segment
        issues, next_issue = setup_summarised_data(segment.segmentIssuesDetected)
        actions, next_action = setup_summarised_data(segment.segmentActionItemsDetected)
        outcomes, next_outcome = setup_summarised_data(segment.segmentOutcomesDetected)

        # Then do each word with confidence-level colouring
        text_index = 1
        live_issue = False
        live_action = False
        live_outcome = False
        for eachWord in segment.segmentConfidence:
            # Look to start a new summary block if needed, in strict priority order - issues, actions, then outcomes.
            # We cannot start a new one until an existing one finishes, so if 2 overlap (unlikely) we skip the second
            live_issue = start_summary_run_highlight(content_col_offset, live_issue, live_action or live_outcome,
                                                     next_issue, row_cells, text_index, "[ISSUE]")
            live_action = start_summary_run_highlight(content_col_offset, live_action, live_issue or live_outcome,
                                                      next_action, row_cells, text_index, "[ACTION]")
            live_outcome = start_summary_run_highlight(content_col_offset, live_outcome, live_issue or live_action,
                                                       next_outcome, row_cells, text_index, "[OUTCOME]")

            # Output the next word, with the correct confidence styling and forced background
            run = row_cells[COL_CONTENT + content_col_offset].paragraphs[0].add_run(eachWord["text"])
            text_index += len(eachWord["text"])
            confLevel = eachWord["confidence"]
            set_transcript_text_style(run, live_issue or live_outcome or live_action, confidence=confLevel)

            # Has any in-progress summarisation block now finished?  Check each one
            live_issue, next_issue = stop_summary_run_highlight(issues, live_issue, next_issue, text_index)
            live_action, next_action = stop_summary_run_highlight(actions, live_action, next_action, text_index)
            live_outcome, next_outcome = stop_summary_run_highlight(outcomes, live_outcome, next_outcome, text_index)

        # If enabled, finish with the base sentiment for the segment - don't write out
        # score if it turns out that this segment ie neither Negative nor Positive
        if sentiment_enabled:
            if segment.segmentIsPositive or segment.segmentIsNegative:
                paragraph = row_cells[COL_SENTIMENT].paragraphs[0]
                img_run = paragraph.add_run()
                if segment.segmentIsPositive:
                    img_run.add_picture(png_smile, width=Mm(4))
                else:
                    img_run.add_picture(png_frown, width=Mm(4))

                # We only have turn-by-turn sentiment score values in non-analytics mode
                if not analytics_mode:
                    text_run = paragraph.add_run(' (' + str(segment.segmentSentimentScore)[:4] + ')')
                    text_run.font.size = Pt(7)
                    text_run.font.italic = True
            else:
                row_cells[COL_SENTIMENT].paragraphs[0].add_run().add_picture(png_neutral, width=Mm(4))

        # Add highlighting to the row if required
        if shading_reqd:
            for column in range(0, COL_CONTENT + content_col_offset + 1):
                set_table_cell_background_colour(row_cells[column], ALTERNATE_ROW_COLOUR)
        shading_reqd = not shading_reqd

        # Check if a category occurs in the middle of a segment - put it after the segment, as timestamp is "later"
        for category_start in keyed_categories.copy().keys():
            if (start_in_millis < category_start) and (category_start < end_in_millis):
                insert_category_row(content_col_offset, keyed_categories, output_table, category_start)
                keyed_categories.pop(category_start)

        # Before we end, does an analytics category start with this line's end time?
        if end_in_millis in keyed_categories:
            # If so, write out the line after this
            insert_category_row(content_col_offset, keyed_categories, output_table, end_in_millis)
            keyed_categories.pop(end_in_millis)


def stop_summary_run_highlight(summaries, live_summary, next_summary, text_index):
    """
    Checks the supplied flags to see that particular type of call summary - e.g. issues or actions - has
    reached the end of it's final word.  If so then it resets the flags and shifts the structures to
    the next summary item of that type in this segment (there most-likely aren't any more)

    :param summaries: List of remaining summary data items to be fully-processed
    :param live_summary: Flag to indicate is this type of call summary data is currently running
    :param next_summary: Start/end word offset information for the current/next summary data item
    :param text_index: Text offset position for this segment what we've rendered up to
    """

    if live_summary and next_summary["End"] <= text_index:
        # Yes - stop highlighting, and pick up any pending summary left on this line of this type
        live_summary = False
        if len(summaries) > 0:
            next_summary = summaries.pop()
        else:
            next_summary = {}
    return live_summary, next_summary


def start_summary_run_highlight(content_col_offset, this_summary, other_summaries, next_summ_item, row_cells,
                                text_index, output_phrase):
    """
    This looks at a call summary data block to see if it has started - if it has then we output a
    message with a highlight and set the text-run highlighting to continue.  If a summary block of
    any other type is currently in-progress then we skip displaying this one, as in a Word document
    the highlighting would be confusing and hard to do.

    :param content_col_offset: Offset into the Word table so we can skip non-existent sentiment columns
    :param this_summary: Flag indicating if a highlighting run for this type is already in progress
    :param other_summaries: Flag indicating if a highlighting run for any other type is already in progress
    :param next_summ_item: The next summary item to be considered for highlighting
    :param row_cells: Cell reference in the Word table for the current speech segment
    :param text_index: Text offset position for this segment what we've rendered up to
    :param output_phrase: Phrase to use in the transcript to mark the start of this highighting run
    """

    new_summary = this_summary

    if len(next_summ_item) > 0 and not this_summary and not other_summaries:
        if (next_summ_item["Begin"] == 0 and text_index == 1) or (next_summ_item["Begin"] == text_index):
            # If so, start the highlighting run, tagging on a leading/trailing
            # highlight space depending on where were are in the segment
            if text_index == 1:
                next_phrase = output_phrase + " "
            else:
                next_phrase = " " + output_phrase
            run = row_cells[COL_CONTENT + content_col_offset].paragraphs[0].add_run(next_phrase)
            set_transcript_text_style(run, True, confidence=0.0)
            new_summary = True

    return new_summary


def setup_summarised_data(summary_block):
    """
    Creates a copy of specified call-summary data block in preparation for writing out the transcription.  This is
    used for each of the supported summary data types.  Returns the first item in the block, or {} if there
    aren't any items, as well as the copy of the block minus the header item

    :param summary_block: The summarise block of data that we're interested in
    """
    summary_data = summary_block.copy()
    if len(summary_data) > 0:
        next_data_item = summary_data.pop()
    else:
        next_data_item = {}
    return summary_data, next_data_item


def insert_category_row(content_col_offset, keyed_categories, output_table, timestamp_millis):
    """
    When writing out the transcript table this method will add in an additional row based
    upon the found entry in the time-keyed category list

    :param content_col_offset: Any additionl
    :param keyed_categories: List of categories identified at any timestamps
    :param output_table: Word document structure to write the table into
    :param timestamp_millis: Timestamp key whose data we have to write out (in milliseconds)
    """

    # Create a new row with the timestamp leading cell, then merge the other cells together
    row_cells = output_table.add_row().cells
    row_cells[COL_STARTTIME].text = convert_timestamp(timestamp_millis / 1000.0)
    merged_cells = row_cells[COL_ENDTIME].merge(row_cells[COL_CONTENT + content_col_offset])

    # Insert the text for each found category
    run = merged_cells.paragraphs[0].add_run("[CATEGORY]")
    set_transcript_text_style(run, False, rgb_color=CATEGORY_TRANSCRIPT_FG_COLOUR)
    run = merged_cells.paragraphs[0].add_run(" " + " ".join(keyed_categories[timestamp_millis]))
    set_transcript_text_style(run, False, confidence=0.5)

    # Give this row a special colour so that it stands out when scrolling
    set_table_cell_background_colour(row_cells[COL_STARTTIME], CATEGORY_TRANSCRIPT_BG_COLOUR)
    set_table_cell_background_colour(merged_cells, CATEGORY_TRANSCRIPT_BG_COLOUR)


def merge_speaker_segments(input_segment_list):
    """
    Merges together consecutive speaker segments unless:
    (a) There is a speaker change, or
    (b) The gap between segments is greater than our acceptable level of delay

    :param input_segment_list: Full time-sorted list of speaker segments
    :return: An updated segment list
    """
    outputSegmentList = []
    lastSpeaker = ""
    lastSegment = None

    # Step through each of our defined speaker segments
    for segment in input_segment_list:
        if (segment.segmentSpeaker != lastSpeaker) or \
                ((segment.segmentStartTime - lastSegment.segmentEndTime) >= START_NEW_SEGMENT_DELAY):
            # Simple case - speaker change or > n-second gap means new output segment
            outputSegmentList.append(segment)

            # This is now our base segment moving forward
            lastSpeaker = segment.segmentSpeaker
            lastSegment = segment
        else:
            # Same speaker, short time, need to copy this info to the last one
            lastSegment.segmentEndTime = segment.segmentEndTime
            lastSegment.segmentText += " " + segment.segmentText
            segment.segmentConfidence[0]["text"] = " " + segment.segmentConfidence[0]["text"]
            for wordConfidence in segment.segmentConfidence:
                lastSegment.segmentConfidence.append(wordConfidence)

    return outputSegmentList


def generate_sentiment(segment_list, language_code):
    """
    Generates sentiment per speech segment, inserting the results into the input list.  This will use
    Amazon Comprehend, but we need to map the job language code to one that Comprehend understands

    :param segment_list: List of speech segments
    :param language_code: Language code to use for the Comprehend job
    """
    # Get our botot3 client, then go through each segment
    client = boto3.client("comprehend")
    for nextSegment in segment_list:
        if len(nextSegment.segmentText) >= MIN_SENTIMENT_LENGTH:
            nextText = nextSegment.segmentText
            response = client.detect_sentiment(Text=nextText, LanguageCode=language_code)
            positiveBase = response["SentimentScore"]["Positive"]
            negativeBase = response["SentimentScore"]["Negative"]

            # If we're over the NEGATIVE threshold then we're negative
            if negativeBase >= MIN_SENTIMENT_NEGATIVE:
                nextSegment.segmentIsNegative = True
                nextSegment.segmentSentimentScore = negativeBase
            # Else if we're over the POSITIVE threshold then we're positive,
            # otherwise we're either MIXED or NEUTRAL and we don't really care
            elif positiveBase >= MIN_SENTIMENT_POSITIVE:
                nextSegment.segmentIsPositive = True
                nextSegment.segmentSentimentScore = positiveBase

            # Store all of the original sentiments for future use
            nextSegment.segmentAllSentiments = response["SentimentScore"]
            nextSegment.segmentPositive = positiveBase
            nextSegment.segmentNegative = negativeBase


def set_repeat_table_header(row):
    """
    Set Word repeat table row on every new page
    """
    row_pointer = row._tr.get_or_add_trPr()
    table_header = OxmlElement('w:tblHeader')
    table_header.set(qn('w:val'), "true")
    row_pointer.append(table_header)
    return row


def load_image(url):
    """
    Loads binary image data from a URL for later embedding into a docx document
    :param url: URL of image to be downloaded
    :return: BytesIO object that can be added as a docx image
    """
    image_url = urllib.request.urlopen(url)
    io_url = BytesIO()
    io_url.write(image_url.read())
    io_url.seek(0)
    return io_url


def write_small_header_text(document, text, confidence):
    """
    Helper function to write out small header entries, where the text colour matches the
    colour of the transcript text for a given confidence value

    :param document: Document to write the text to
    :param text: Text to be output
    :param confidence: Confidence score, which changes the text colour
    """
    run = document.paragraphs[-1].add_run(text)
    set_transcript_text_style(run, False, confidence=confidence)
    run.font.size = Pt(7)
    run.font.italic = True


def write(cli_arguments, speech_segments, job_status, summaries_detected):
    """
    Write a transcript from the .json transcription file and other data generated
    by the results parser, putting it all into a human-readable Word document

    :param cli_arguments: CLI arguments used for this processing run
    :param speech_segments: List of call speech segments
    :param job_status: Status of the Transcribe job
    :param summaries_detected: Flag to indicate presence of call summary data
    """

    json_filepath = Path(cli_arguments.inputFile)
    data = json.load(open(json_filepath.absolute(), "r", encoding="utf-8"))
    sentimentEnabled = (cli_arguments.sentiment == 'on')
    tempFiles = []

    # Initiate Document, orientation and margins
    document = Document()
    document.sections[0].left_margin = Mm(19.1)
    document.sections[0].right_margin = Mm(19.1)
    document.sections[0].top_margin = Mm(19.1)
    document.sections[0].bottom_margin = Mm(19.1)
    document.sections[0].page_width = Mm(210)
    document.sections[0].page_height = Mm(297)

    # Set the base font and document title
    font = document.styles["Normal"].font
    font.name = "Calibri"
    font.size = Pt(10)

    # Create our custom text header style
    custom_style = document.styles.add_style(CUSTOM_STYLE_HEADER, WD_STYLE_TYPE.PARAGRAPH)
    custom_style.paragraph_format.widow_control = True
    custom_style.paragraph_format.keep_with_next = True
    custom_style.paragraph_format.space_after = Pt(0)
    custom_style.font.size = font.size
    custom_style.font.name = font.name
    custom_style.font.bold = True
    custom_style.font.italic = True

    # Intro banner header
    document.add_picture(load_image(IMAGE_URL_BANNER), width=Mm(171))

    # Pull out header information - some from the JSON, but most only exists in the Transcribe job status
    if cli_arguments.analyticsMode:
        # We need 2 columns only if we're in analytics mode, as we put the charts on the right of the table
        document.add_section(WD_SECTION.CONTINUOUS)
        section_ptr = document.sections[-1]._sectPr
        cols = section_ptr.xpath('./w:cols')[0]
        cols.set(qn('w:num'), '2')

    # Write put the call summary table - depending on the mode that Transcribe was used in, and
    # if the request is being run on a JSON results file rather than reading the job info from Transcribe,
    # not all of the information is available.
    # -- Media information
    # -- Amazon Transcribe job information
    # -- Average transcript word-confidence scores
    write_custom_text_header(document, "Amazon Transcribe Audio Source")
    table = document.add_table(rows=1, cols=2)
    table.style = document.styles[TABLE_STYLE_STANDARD]
    table.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Job Name"
    if cli_arguments.analyticsMode:
        hdr_cells[1].text = data["JobName"]
    else:
        hdr_cells[1].text = data["jobName"]
    job_data = []
    # Audio duration is the end-time of the final voice segment, which might be shorter than the actual file duration
    if len(speech_segments) > 0:
        audio_duration = speech_segments[-1].segmentEndTime
        dur_text = str(int(audio_duration / 60)) + "m " + str(round(audio_duration % 60, 2)) + "s"
        job_data.append({"name": "Audio Duration", "value": dur_text})
    # We can infer diarization mode from the JSON results data structure
    if cli_arguments.analyticsMode:
        job_data.append({"name": "Audio Ident", "value": "Call Analytics"})
    elif "speaker_labels" in data["results"]:
        job_data.append({"name": "Audio Ident", "value": "Speaker-separated"})
    else:
        job_data.append({"name": "Audio Ident", "value": "Channel-separated"})

    # Some information is only in the job status
    if job_status is not None:
        job_data.append({"name": "Language", "value": job_status["LanguageCode"]})
        job_data.append({"name": "File Format", "value": job_status["MediaFormat"]})
        job_data.append({"name": "Sample Rate", "value": str(job_status["MediaSampleRateHertz"]) + " Hz"})
        job_data.append({"name": "Job Created", "value": job_status["CreationTime"].strftime("%a %d %b '%y at %X")})
        if "ContentRedaction" in job_status["Settings"]:
            redact_type = job_status["Settings"]["ContentRedaction"]["RedactionType"]
            redact_output = job_status["Settings"]["ContentRedaction"]["RedactionOutput"]
            job_data.append({"name": "Redaction Mode", "value": redact_type + " [" + redact_output + "]"})
        if "VocabularyFilterName" in job_status["Settings"]:
            vocab_filter = job_status["Settings"]["VocabularyFilterName"]
            vocab_method = job_status["Settings"]["VocabularyFilterMethod"]
            job_data.append({"name": "Vocabulary Filter", "value": vocab_filter + " [" + vocab_method + "]"})
        if "VocabularyName" in job_status["Settings"]:
            job_data.append({"name": "Custom Vocabulary", "value": job_status["Settings"]["VocabularyName"]})

    # Finish with the confidence scores (if we have any)
    stats = generate_confidence_stats(speech_segments)
    if len(stats["accuracy"]) > 0:
        job_data.append({"name": "Avg. Confidence", "value": str(round(statistics.mean(stats["accuracy"]), 2)) + "%"})

    # Place all of our job-summary fields into the Table, one row at a time
    for next_row in job_data:
        row_cells = table.add_row().cells
        row_cells[0].text = next_row["name"]
        row_cells[1].text = next_row["value"]

    # Formatting transcript table widths
    widths = (Cm(3.44), Cm(4.89))
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

    # Spacer paragraph
    document.add_paragraph()

    # Conversational Analytics (other column) if enabled
    # -- Caller sentiment graph
    # -- Talk time split
    if cli_arguments.analyticsMode:
        write_header_graphs(data, document, tempFiles)

    # At this point, if we have no transcript then we need to quickly exit
    if len(speech_segments) == 0:
        document.add_section(WD_SECTION.CONTINUOUS)
        section_ptr = document.sections[-1]._sectPr
        cols = section_ptr.xpath('./w:cols')[0]
        cols.set(qn('w:num'), '1')
        write_custom_text_header(document, "This call had no audible speech to transcribe.")
    else:
        # Conversational Analytics (new Section)
        # -- Show speaker loudness graph, with sentiment, interrupts and non-talk time highlighted
        # -- Show a summary of any call analytics categories detected
        # -- Show a summary of any issues detected in the transcript
        # -- Process and display speaker sentiment by period
        if cli_arguments.analyticsMode:
            build_call_loudness_charts(document, speech_segments, data["ConversationCharacteristics"]["Interruptions"],
                                       data["ConversationCharacteristics"]["NonTalkTime"],
                                       data["ConversationCharacteristics"]["TalkTime"], tempFiles)
            keyed_categories = write_detected_categories(document, data["Categories"]["MatchedDetails"])
            write_analytics_sentiment(data, document)

            # Write out any call summarisation data
            if summaries_detected:
                write_detected_summaries(document, speech_segments)
        else:
            # No analytics => no categories
            keyed_categories = {}

        # Process and display transcript by speaker segments (new section)
        # -- Conversation "turn" start time and duration
        # -- Speaker identification
        # -- Sentiment type (if enabled) and sentiment score (if available)
        # -- Transcribed text with (if available) Call Analytics markers
        document.add_section(WD_SECTION.CONTINUOUS)
        section_ptr = document.sections[-1]._sectPr
        cols = section_ptr.xpath('./w:cols')[0]
        cols.set(qn('w:num'), '1')
        write_custom_text_header(document, "Call Transcription")
        document.add_paragraph()  # Spacing
        write_small_header_text(document, "WORD CONFIDENCE: >= 90% in black, ", 0.9)
        write_small_header_text(document, ">= 50% in brown, ", 0.5)
        write_small_header_text(document, "< 50% in red", 0.49)
        table_cols = 4
        if sentimentEnabled or cli_arguments.analyticsMode:
            # Ensure that we add space for the sentiment column
            table_cols += 1
            content_col_offset = 0
        else:
            # Will need to shift the content column to the left, as Sentiment isn't there now
            content_col_offset = -1
        table = document.add_table(rows=1, cols=table_cols)
        table.style = document.styles[TABLE_STYLE_STANDARD]
        hdr_cells = table.rows[0].cells
        hdr_cells[COL_STARTTIME].text = "Start"
        hdr_cells[COL_ENDTIME].text = "Dur."
        hdr_cells[COL_SPEAKER].text = "Speaker"
        hdr_cells[COL_CONTENT + content_col_offset].text = "Transcription"

        # Based upon our segment list, write out the transcription table
        write_transcribe_text(table, sentimentEnabled or cli_arguments.analyticsMode, cli_arguments.analyticsMode,
                              speech_segments, keyed_categories)
        document.add_paragraph()

        # Formatting transcript table widths - we need to add sentiment
        # column if needed, and it and the content width accordingly
        widths = [Inches(0.8), Inches(0.5), Inches(0.5), 0]
        if sentimentEnabled:
            # Comprehend sentiment needs space for the icon and % score
            widths.append(0)
            widths[COL_CONTENT + + content_col_offset] = Inches(7)
            widths[COL_SENTIMENT] = Inches(0.7)
        elif cli_arguments.analyticsMode:
            # Analytics sentiment just needs an icon
            widths.append(0)
            widths[COL_CONTENT + + content_col_offset] = Inches(7.4)
            widths[COL_SENTIMENT] = Inches(0.3)
        else:
            widths[COL_CONTENT + content_col_offset] = Inches(7.7)
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width

        # Setup the repeating header
        set_repeat_table_header(table.rows[0])

        # Display confidence count table, if requested (new section)
        # -- Summary table of confidence scores into "bins"
        # -- Scatter plot of confidence scores over the whole transcript
        if cli_arguments.confidence == 'on':
            write_confidence_scores(document, stats, tempFiles)
            document.add_section(WD_SECTION.CONTINUOUS)

        # Generate our raw data for the Comprehend sentiment graph (if requested)
        if sentimentEnabled:
            write_comprehend_sentiment(document, speech_segments, tempFiles)

    # Save the whole document
    document.save(cli_arguments.outputFile)

    # Now delete any local images that we created
    for filename in tempFiles:
        os.remove(filename)


def write_header_graphs(data, document, temp_files):
    """
    Writes out the two header-level graphs for caller sentiment and talk-time split

    :param data: JSON result data from Transcribe
    :param document: Word document structure to write the table into
    :param temp_files: List of temporary files for later deletion
    """
    characteristics = data["ConversationCharacteristics"]
    # Caller sentiment graph
    fig, ax = plt.subplots(nrows=1, ncols=2, figsize=(12.5 / 2.54, 8 / 2.54), gridspec_kw={'width_ratios': [4, 3]})
    period_sentiment = characteristics["Sentiment"]["SentimentByPeriod"]["QUARTER"]
    # Graph configuration
    ax[0].set_xlim(xmin=1, xmax=4)
    ax[0].set_ylim(ymax=5, ymin=-5)
    ax[0].yaxis.set_major_locator(ticker.MultipleLocator(5.0))
    ax[0].spines['bottom'].set_position('zero')
    ax[0].spines['top'].set_color('none')
    ax[0].spines['right'].set_color('none')
    ax[0].set_xticks([])
    ax[0].set_title("Customer sentiment", fontsize=10, fontweight="bold", pad="12.0")
    # Only draw the sentiment line if we actually have a Customer that talked
    if "CUSTOMER" in period_sentiment:
        # Setup our data holders, then extract it all
        x_sentiment = np.array([])
        y_sentiment = np.array([])
        period_index = 1
        for score in period_sentiment["CUSTOMER"]:
            x_sentiment = np.append(x_sentiment, period_index)
            y_sentiment = np.append(y_sentiment, score["Score"])
            period_index += 1

        # Set the line colour to match the overall sentiment
        if characteristics["Sentiment"]["OverallSentiment"]["CUSTOMER"] >= 0.0:
            line_colour = "darkgreen"
        else:
            line_colour = "red"

        # Now draw out the simple line plot
        x_new = np.linspace(1, 4, 200)
        spline = make_interp_spline(x_sentiment, y_sentiment)
        y_smooth = spline(x_new)
        ax[0].plot(x_new, y_smooth, linewidth=3, color=line_colour)
    # Talk time calculations and ratios
    non_talk = characteristics["NonTalkTime"]["Instances"]
    quiet_time = 0
    for quiet in non_talk:
        quiet_time += quiet["DurationMillis"]
    if "AGENT" in characteristics["TalkTime"]["DetailsByParticipant"]:
        agent_talk_time = characteristics["TalkTime"]["DetailsByParticipant"]["AGENT"]["TotalTimeMillis"]
    else:
        agent_talk_time = 0
    if "CUSTOMER" in characteristics["TalkTime"]["DetailsByParticipant"]:
        caller_talk_time = characteristics["TalkTime"]["DetailsByParticipant"]["CUSTOMER"]["TotalTimeMillis"]
    else:
        caller_talk_time = 0
    total_time = agent_talk_time + caller_talk_time + quiet_time
    if total_time > 0:
        quiet_ratio = quiet_time / total_time * 100.0
        agent_ratio = agent_talk_time / total_time * 100.0
        caller_ratio = caller_talk_time / total_time * 100.0
    else:
        quiet_ratio = 0.0
        agent_ratio = 0.0
        caller_ratio = 0.0
    ratio_format = "{speaker} ({ratio:.1f}%)"
    # Additional configuration
    ax[1].set_xticks([])
    ax[1].set_yticks([])
    ax[1].set_title("Talk time", fontsize=10, fontweight="bold", pad="10.0")
    ax[1].spines['top'].set_color('none')
    ax[1].spines['bottom'].set_color('none')
    ax[1].spines['left'].set_color('none')
    ax[1].spines['right'].set_color('none')
    # Now draw out the plot
    labels = ["time"]
    width = 1.0
    ax[1].bar(labels, [quiet_time], width, label=ratio_format.format(ratio=quiet_ratio, speaker="Non-Talk"),
              bottom=[agent_talk_time + caller_talk_time])
    ax[1].bar(labels, [caller_talk_time], width, label=ratio_format.format(ratio=caller_ratio, speaker="Customer"),
              bottom=[agent_talk_time])
    ax[1].bar(labels, [agent_talk_time], width, label=ratio_format.format(ratio=agent_ratio, speaker="Agent"))
    box = ax[1].get_position()
    ax[1].set_position([box.x0, box.y0 + box.height * 0.25, box.width, box.height * 0.75])
    ax[1].legend(loc="upper center", bbox_to_anchor=(0.5, -0.05), ncol=1)
    chart_file_name = "./" + "talk-time.png"
    plt.savefig(chart_file_name, facecolor="aliceblue")
    temp_files.append(chart_file_name)
    document.add_picture(chart_file_name, width=Cm(7.5))
    plt.clf()


def generate_confidence_stats(speech_segments):
    """
    Creates a map of timestamps and confidence scores to allow for both summarising and graphing in the document.
    We also need to bucket the stats for summarising into bucket ranges that feel important (but are easily changed)
    
    :param speech_segments: List of call speech segments 
    :return: Confidence and timestamp structures for graphing 
    """""

    # Stats dictionary
    stats = {
        "timestamps": [],
        "accuracy": [],
        "9.8": 0, "9": 0, "8": 0, "7": 0, "6": 0, "5": 0, "4": 0, "3": 0, "2": 0, "1": 0, "0": 0,
        "parsedWords": 0}

    # Confidence count - we need the average confidence score regardless
    for line in speech_segments:
        for word in line.segmentConfidence:
            stats["timestamps"].append(word["start_time"])
            conf_value = word["confidence"]
            stats["accuracy"].append(int(conf_value * 100))
            if conf_value >= 0.98:
                stats["9.8"] += 1
            elif conf_value >= 0.9:
                stats["9"] += 1
            elif conf_value >= 0.8:
                stats["8"] += 1
            elif conf_value >= 0.7:
                stats["7"] += 1
            elif conf_value >= 0.6:
                stats["6"] += 1
            elif conf_value >= 0.5:
                stats["5"] += 1
            elif conf_value >= 0.4:
                stats["4"] += 1
            elif conf_value >= 0.3:
                stats["3"] += 1
            elif conf_value >= 0.2:
                stats["2"] += 1
            elif conf_value >= 0.1:
                stats["1"] += 1
            else:
                stats["0"] += 1
            stats["parsedWords"] += 1
    return stats


def write_custom_text_header(document, text_label):
    """
    Adds a run of text to the document with the given text label, but using our customer text-header style

    :param document: Word document structure to write the table into
    :param text_label: Header text to write out
    :return:
    """
    paragraph = document.add_paragraph(text_label)
    paragraph.style = CUSTOM_STYLE_HEADER


def write_confidence_scores(document, stats, temp_files):
    """
    Using the pre-build confidence stats list, create a summary table of confidence score
    spreads, as well as a scatter-plot showing each word against the overall mean

    :param document: Word document structure to write the table into
    :param stats: Statistics for the confidence scores in the conversation
    :param temp_files: List of temporary files for later deletion
    :return:
    """
    document.add_section(WD_SECTION.CONTINUOUS)
    section_ptr = document.sections[-1]._sectPr
    cols = section_ptr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), '2')
    write_custom_text_header(document, "Word Confidence Scores")
    # Start with the fixed headers
    table = document.add_table(rows=1, cols=3)
    table.style = document.styles[TABLE_STYLE_STANDARD]
    table.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Confidence"
    hdr_cells[1].text = "Count"
    hdr_cells[2].text = "Percentage"
    parsedWords = stats["parsedWords"]
    confidenceRanges = ["98% - 100%", "90% - 97%", "80% - 89%", "70% - 79%", "60% - 69%", "50% - 59%", "40% - 49%",
                        "30% - 39%", "20% - 29%", "10% - 19%", "0% - 9%"]
    confidenceRangeStats = ["9.8", "9", "8", "7", "6", "5", "4", "3", "2", "1", "0"]
    # Add on each row
    shading_reqd = False
    for confRange, rangeStats in zip(confidenceRanges, confidenceRangeStats):
        row_cells = table.add_row().cells
        row_cells[0].text = confRange
        row_cells[1].text = str(stats[rangeStats])
        row_cells[2].text = str(round(stats[rangeStats] / parsedWords * 100, 2)) + "%"

        # Add highlighting to the row if required
        if shading_reqd:
            for column in range(0, 3):
                set_table_cell_background_colour(row_cells[column], ALTERNATE_ROW_COLOUR)
        shading_reqd = not shading_reqd

    # Formatting transcript table widths, then move to the next column
    widths = (Inches(1.2), Inches(0.8), Inches(0.8))
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width
    # Confidence of each word as scatter graph, and the mean as a line across
    fig, ax = plt.subplots(nrows=1, ncols=1, figsize=(6, 4))
    ax.scatter(stats["timestamps"], stats["accuracy"])
    ax.plot([stats["timestamps"][0], stats["timestamps"][-1]], [statistics.mean(stats["accuracy"]),
                                                                statistics.mean(stats["accuracy"])], "r")
    # Formatting
    ax.set_xlabel("Time (seconds)")
    ax.set_ylabel("Word Confidence (percent)")
    ax.set_yticks(range(0, 101, 10))
    fig.suptitle("Word Confidence During Transcription", fontsize=11, fontweight="bold")
    ax.legend(["Word Confidence Mean", "Individual words"], loc="lower center")
    # Write out the chart
    chart_file_name = "./" + "chart.png"
    plt.savefig(chart_file_name, facecolor="aliceblue")
    temp_files.append(chart_file_name)
    plt.clf()
    document.add_picture(chart_file_name, width=Cm(8))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.LEFT
    document.add_paragraph()


def insert_line_and_col_break(document):
    """
    Inserts a line break and column break into the document

    :param document: Word document structure to write the breaks into
    """
    # Blank line followed by column break
    document.add_paragraph()  # Spacing
    run = document.paragraphs[-1].add_run()
    run.add_break(WD_BREAK.LINE)
    run.add_break(WD_BREAK.COLUMN)


def write_detected_categories(document, category_list):
    """
    If there are any detected categories then write out a simple list

    :param document: Word document structure to write the table into
    :param category_list: Details of detected categories
    :return: A timestamp-keyed list of detected categories, which we'll use later when writing out the transcript
    """
    timed_categories = {}
    if category_list != {}:
        # Start with a new single-column section
        document.add_section(WD_SECTION.CONTINUOUS)
        section_ptr = document.sections[-1]._sectPr
        cols = section_ptr.xpath('./w:cols')[0]
        cols.set(qn('w:num'), '1')
        write_custom_text_header(document, "Categories Detected")

        # Table header information
        table = document.add_table(rows=1, cols=3)
        table.style = document.styles[TABLE_STYLE_STANDARD]
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Category"
        hdr_cells[1].text = "#"
        hdr_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[2].text = "Timestamps found at"

        # Go through each detected category
        for next_cat in category_list.keys():
            row_cells = table.add_row().cells
            row_cells[0].text = next_cat

            # Instances and timestamps for the category do not exist for "negative" categories
            if category_list[next_cat]["PointsOfInterest"] != []:
                row_cells[1].text = str(len(category_list[next_cat]["PointsOfInterest"]))
                row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Now go through each instance of it
                instance = 0
                for next_timestamp in category_list[next_cat]["PointsOfInterest"]:
                    # Add the next timestamp to the document row, with separating punctuation if needed
                    start_time_millis = next_timestamp["BeginOffsetMillis"]
                    start_time_text = convert_timestamp(start_time_millis / 1000.0)
                    if instance > 0:
                        row_cells[2].paragraphs[0].add_run(", ")
                    row_cells[2].paragraphs[0].add_run(start_time_text)
                    instance += 1

                    # Now add this to our time-keyed category list
                    if start_time_millis not in timed_categories:
                        timed_categories[start_time_millis] = [next_cat]
                    else:
                        timed_categories[start_time_millis].append(next_cat)

        # Formatting transcript table widths
        widths = (Cm(4.0), Cm(1.0), Cm(12.2))
        shading_reqd = False
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width
                if shading_reqd:
                    set_table_cell_background_colour(row.cells[idx], ALTERNATE_ROW_COLOUR)
            shading_reqd = not shading_reqd

        # Finish with some spacing
        document.add_paragraph()

    # Return our time-keyed category list
    return timed_categories


def write_detected_summaries(document, speech_segments):
    """
    Scans the speech segments for any detected summaries of the requested type, and if there are any
    then a new table is added to the document.  This assumes that we do have some summaries, as if
    not we'll just output a table header on its own

    :param document: Word document structure to write the table into
    :param speech_segments: Call transcript structures
    """

    # Start with a new single-column section
    document.add_section(WD_SECTION.CONTINUOUS)
    section_ptr = document.sections[-1]._sectPr
    cols = section_ptr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), '1')
    table = document.add_table(rows=1, cols=3)
    table.style = document.styles[TABLE_STYLE_STANDARD]
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Call Summary Highlights"
    hdr_cells[0].merge(hdr_cells[2])

    # Loop through each of our summary types
    for summary_map in CALL_SUMMARY_MAP:
        # Scan through the segments and extract the issues
        summary_detected = []
        for turn in speech_segments:
            summary_block = getattr(turn, summary_map["Field"])
            # for issue in turn.myVar:
            for issue in summary_block:
                new_summary = {"Speaker": turn.segmentSpeaker}
                new_summary["Timestamp"] = turn.segmentStartTime
                new_summary["Text"] = turn.segmentText[issue["Begin"]:issue["End"]]
                # May need a prefix or suffix for partial text
                if issue["Begin"] > 0:
                    new_summary["Text"] = "..." + new_summary["Text"]
                if issue["End"] < len(turn.segmentText):
                    new_summary["Text"] = new_summary["Text"] + "..."
                summary_detected.append(new_summary)

        # If we found some of this type then write out a table
        if summary_detected:
            # Header section for this block
            row_cells = table.add_row().cells
            row_cells[0].text = summary_map["Title"]
            set_table_cell_background_colour(row_cells[0], summary_map["Color"])
            row_cells[0].merge(row_cells[2])

            # Column header section for this block
            next_row = table.add_row()
            row_cells = next_row.cells
            row_cells[0].text = "Speaker"
            row_cells[1].text = "Turn Time"
            row_cells[2].text = "Detected Text"
            set_table_row_bold(next_row, True)
            shading_reqd = False

            # Output each row
            for issue in summary_detected:
                # First column is the speaker
                next_row = table.add_row()
                row_cells = next_row.cells
                row_cells[0].text = issue["Speaker"]
                row_cells[1].text = convert_timestamp(issue["Timestamp"])
                row_cells[2].text = issue["Text"]
                set_table_row_bold(next_row, False)

                # Add highlighting to the row if required; e.g. every 2nd row
                if shading_reqd:
                    for column in range(0, 3):
                        set_table_cell_background_colour(row_cells[column], ALTERNATE_ROW_COLOUR)
                shading_reqd = not shading_reqd

            # Formatting transcript table widths
            widths = (Cm(2.2), Cm(2.2), Cm(12.8))
            for row in table.rows:
                for idx, width in enumerate(widths):
                    row.cells[idx].width = width

    # Finish with some spacing
    document.add_paragraph()


def build_call_loudness_charts(document, speech_segments, interruptions, quiet_time, talk_time, temp_files):
    """
    Creates the call loudness charts for each caller, which we also overlay sentiment on
    :param document: Word document structure to write the graphics into
    :param speech_segments: Call transcript structures
    :param interruptions: Call speaker interruption structures
    :param quiet_time: Call non-talk time structures
    :param talk_time: Call talk time structures
    :param temp_files: List of temporary files for later deletion (includes our graph)
    """

    # Start with a new single-column section
    document.add_section(WD_SECTION.CONTINUOUS)
    section_ptr = document.sections[-1]._sectPr
    cols = section_ptr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), '1')
    document.add_paragraph()
    write_custom_text_header(document, "Conversation Volume Levels with Sentiment and Interruptions")

    # Initialise our loudness structures
    secsLoudAgent = []
    dbLoudAgent = []
    secsLoudCaller = []
    dbLoudCaller = []

    # Work through each conversation turn, extracting timestamp/decibel values as we go
    for segment in speech_segments:
        this_second = int(segment.segmentStartTime)
        # Each segment has a loudness score per second or part second
        for score in segment.segmentLoudnessScores:
            # This can be set to NONE, which causes errors later
            if score is None:
                score = 0.0
            # Track the Agent loudness
            if segment.segmentSpeaker == "Agent":
                secsLoudAgent.append(this_second)
                dbLoudAgent.append(score)
            # Track the Caller loudness
            else:
                secsLoudCaller.append(this_second)
                dbLoudCaller.append(score)
            this_second += 1
    agentLoudness = {"Seconds": secsLoudAgent, "dB": dbLoudAgent}
    callerLoudness = {"Seconds": secsLoudCaller, "dB": dbLoudCaller}

    # Work out our final talk "second", as we need both charts to line up, but
    # be careful as there may just be one speaker in the Call Analytics output
    if talk_time["DetailsByParticipant"]["AGENT"]["TotalTimeMillis"] == 0:
        final_second = max(secsLoudCaller)
        max_decibel = max(dbLoudCaller)
        haveAgent = False
        haveCaller = True
        plotRows = 1
    elif talk_time["DetailsByParticipant"]["CUSTOMER"]["TotalTimeMillis"] == 0:
        final_second = max(secsLoudAgent)
        max_decibel = max(dbLoudAgent)
        haveAgent = True
        haveCaller = False
        plotRows = 1
    else:
        final_second = max(max(secsLoudAgent), max(secsLoudCaller))
        max_decibel = max(max(dbLoudAgent), max(dbLoudCaller))
        haveAgent = True
        haveCaller = True
        plotRows = 2

    # Add some headroom to our decibel limit to give space for "interruption" markers
    max_decibel_headroom = (int(max_decibel / 10) + 2) * 10

    # Create a dataset for interruptions, which needs to be in the background on both charts
    intSecs = []
    intDb = []
    for speaker in interruptions["InterruptionsByInterrupter"]:
        for entry in interruptions["InterruptionsByInterrupter"][speaker]:
            start = int(entry["BeginOffsetMillis"] / 1000)
            end = int(entry["EndOffsetMillis"] / 1000)
            for second in range(start, end+1):
                intSecs.append(second)
                intDb.append(max_decibel_headroom)
    intSegments = {"Seconds": intSecs, "dB": intDb}

    # Create a dataset for non-talk time, which needs to be in the background on both charts
    quietSecs = []
    quietdB = []
    for quiet_period in quiet_time["Instances"]:
        start = int(quiet_period["BeginOffsetMillis"] / 1000)
        end = int(quiet_period["EndOffsetMillis"] / 1000)
        for second in range(start, end + 1):
            quietSecs.append(second)
            quietdB.append(max_decibel_headroom)
    quietSegments = {"Seconds": quietSecs, "dB": quietdB}

    # Either speaker may be missing, so we cannot assume this is a 2-row or 1-row plot
    # We want a 2-row figure, one row per speaker, but with the interruptions on the background
    fig, ax = plt.subplots(nrows=plotRows, ncols=1, figsize=(12, 2.5 * plotRows))
    if haveAgent:
        if haveCaller:
            build_single_loudness_chart(ax[0], agentLoudness, intSegments, quietSegments, speech_segments,
                                        final_second, max_decibel_headroom, "Agent", False, True)
            build_single_loudness_chart(ax[1], callerLoudness, intSegments, quietSegments, speech_segments,
                                        final_second, max_decibel_headroom, "Customer", True, False)
        else:
            build_single_loudness_chart(ax, agentLoudness, intSegments, quietSegments, speech_segments,
                                        final_second, max_decibel_headroom, "Agent", True, True)
    elif haveCaller:
        build_single_loudness_chart(ax, callerLoudness, intSegments, quietSegments, speech_segments,
                                    final_second, max_decibel_headroom, "Customer", True, True)

    # Add the chart to our document
    chart_file_name = "./" + "volume.png"
    fig.savefig(chart_file_name, facecolor="aliceblue")
    temp_files.append(chart_file_name)
    document.add_picture(chart_file_name, width=Cm(17))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.LEFT
    plt.clf()


def build_single_loudness_chart(axes, loudness, interrupts, quiet_time, speech_segments, xaxis_max, yaxis_max, caller,
                                show_x_legend, show_chart_legend):
    """
    Builds a single loundness/sentiment chart using the given data

    :param axes: Axis to use for the chart in our larger table
    :param loudness: Data series for the speakers loudness levels
    :param interrupts: Data series for marking interrupts on the chart
    :param quiet_time: Data series for marking non-talk time on the chart
    :param speech_segments: Call transcript structures
    :param xaxis_max: Second for the last speech entry in the call, which may not have been this speaker
    :param yaxis_max: Max decibel level in the call, which may not have been this speaker
    :param caller: Name of the caller to check for in the transcript
    :param show_x_legend: Flag to show/hide the x-axis legend
    :param show_chart_legend: Flag to show/hide the top-right graph legend
    """

    # Draw the main loudness data bar-chart
    seconds = loudness["Seconds"]
    decibels = loudness["dB"]
    axes.bar(seconds, decibels, label="Speaker volume", width=BAR_CHART_WIDTH)
    axes.set_xlim(xmin=0, xmax=xaxis_max)
    axes.set_ylim(ymax=yaxis_max)
    if show_x_legend:
        axes.set_xlabel("Time (in seconds)")
    axes.set_ylabel("decibels")

    # Build up sentiment data series for positive and negative, plotting it at the bottom
    x = np.linspace(0, max(seconds), endpoint=True, num=(max(seconds) + 1))
    ypos = np.linspace(0, 0, endpoint=True, num=(max(seconds) + 1))
    yneg = np.linspace(0, 0, endpoint=True, num=(max(seconds) + 1))
    yneut = np.linspace(0, 0, endpoint=True, num=(max(seconds) + 1))
    for segment in speech_segments:
        this_second = int(segment.segmentStartTime)
        if segment.segmentSpeaker == caller:
            if segment.segmentIsPositive:
                for score in segment.segmentLoudnessScores:
                    ypos[this_second] = 10
                    this_second += 1
            elif segment.segmentNegative:
                for score in segment.segmentLoudnessScores:
                    yneg[this_second] = 10
                    this_second += 1
            else:
                for score in segment.segmentLoudnessScores:
                    yneut[this_second] = 10
                    this_second += 1
    axes.bar(x, ypos, label="Positive sentiment", color="limegreen", width=BAR_CHART_WIDTH)
    axes.bar(x, yneg, label="Negative sentiment", color="orangered", width=BAR_CHART_WIDTH)
    axes.bar(x, yneut, label="Neutral sentiment", color="cadetblue", width=BAR_CHART_WIDTH)

    # Finish with the non-talk and interrupt overlays (if there are any)
    if len(quiet_time["Seconds"]) > 0:
        axes.bar(quiet_time["Seconds"], quiet_time["dB"], label="Non-talk time", color="lightcyan", width=BAR_CHART_WIDTH)
    if len(interrupts["Seconds"]) > 0:
        axes.bar(interrupts["Seconds"], interrupts["dB"], label="Interruptions", color="goldenrod", width=BAR_CHART_WIDTH, alpha=0.5, bottom=10)

    # Only show the legend for the top graph if requested
    box = axes.get_position()
    axes.set_position([0.055, box.y0, box.width, box.height])
    axes.text(5, yaxis_max-5, caller, style='normal', color='black', bbox={'facecolor': 'white', 'pad': 5})
    if show_chart_legend:
        axes.legend(loc="upper right", bbox_to_anchor=(1.21, 1.0), ncol=1, borderaxespad=0)


def write_comprehend_sentiment(document, speech_segments, temp_files):
    """
    Writes out tables for per-period, per-speaker sentiment from the analytics mode, as well as
    the overall sentiment for a speaker

    :param document: Docx document to add the sentiment graph to
    :param speech_segments: Process transcript text holding turn-by-turn sentiment
    :param temp_files: List of temp files to be deleted later
    :return:
    """
    # Initialise our base structures
    speaker0labels = ['ch_0', 'spk_0']
    speaker1labels = ['ch_1', 'spk_1']
    speaker0timestamps = []
    speaker0data = []
    speaker1timestamps = []
    speaker1data = []

    # Start with some spacing and a new sub-header
    document.add_paragraph()
    write_custom_text_header(document, "Amazon Comprehend Sentiment")
    # Now step through and process each speech segment's sentiment
    for segment in speech_segments:
        if segment.segmentIsPositive or segment.segmentIsNegative:
            # Only interested in actual sentiment entries
            score = segment.segmentSentimentScore
            timestamp = segment.segmentStartTime

            # Positive re-calculation
            if segment.segmentIsPositive:
                score = 2 * ((1 - (1 - score) / (1 - MIN_SENTIMENT_POSITIVE)) * 0.5)
            # Negative re-calculation
            else:
                score = 2 * ((1 - score) / (1 - MIN_SENTIMENT_NEGATIVE) * 0.5 - 0.5)

            if segment.segmentSpeaker in speaker1labels:
                speaker1data.append(score)
                speaker1timestamps.append(timestamp)
            elif segment.segmentSpeaker in speaker0labels:
                speaker0data.append(score)
                speaker0timestamps.append(timestamp)

    # Spline fit needs at least 4 points for k=3, but 5 works better
    speaker1k = 3
    speaker0k = 3
    if len(speaker1data) < 5:
        speaker1k = 1
    if len(speaker0data) < 5:
        speaker0k = 1

    # Create Speaker-0 graph
    plt.figure(figsize=(8, 5))
    speaker0xnew = np.linspace(speaker0timestamps[0], speaker0timestamps[-1],
                               int((speaker0timestamps[-1] - speaker0timestamps[0]) + 1.0))
    speaker0spl = make_interp_spline(speaker0timestamps, speaker0data, k=speaker0k)
    speaker0powerSmooth = speaker0spl(speaker0xnew)
    plt.plot(speaker0timestamps, speaker0data, "ro")
    plt.plot(speaker0xnew, speaker0powerSmooth, "r", label="Speaker 1")

    # Create Speaker-1 graph
    speaker1xnew = np.linspace(speaker1timestamps[0], speaker1timestamps[-1],
                               int((speaker1timestamps[-1] - speaker1timestamps[0]) + 1.0))
    speaker1spl = make_interp_spline(speaker1timestamps, speaker1data, k=speaker1k)
    speaker1powerSmooth = speaker1spl(speaker1xnew)
    plt.plot(speaker1timestamps, speaker1data, "bo")
    plt.plot(speaker1xnew, speaker1powerSmooth, "b", label="Speaker 2")

    # Draw it out
    plt.title("Call Sentiment - Pos/Neg Only")
    plt.xlabel("Time (seconds)")
    plt.axis([0, max(speaker0timestamps[-1], speaker1timestamps[-1]), -1.5, 1.5])
    plt.legend()
    plt.axhline(y=0, color='k')
    plt.axvline(x=0, color='k')
    plt.grid(True)
    plt.xticks(np.arange(0, max(speaker0timestamps[-1], speaker1timestamps[-1]), 60))
    plt.yticks(np.arange(-1, 1.01, 0.25))

    # Write out the chart
    chart_file_name = "./" + "sentiment.png"
    plt.savefig(chart_file_name)
    temp_files.append(chart_file_name)
    plt.clf()
    document.add_picture(chart_file_name, width=Cm(14.64))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.LEFT


def set_table_cell_background_colour(cell, rgb_hex):
    """
    Modifies the background color of the given table cell to the given RGB hex value.  This currently isn't
    supporting by the DOCX module, and the only option is to modify the underlying Word document XML

    :param cell: Table cell to be changed
    :param rgb_hex: RBG hex string for the background color
    """
    parsed_xml = parse_xml(r'<w:shd {0} w:fill="{1}"/>'.format(nsdecls('w'), rgb_hex))
    cell._tc.get_or_add_tcPr().append(parsed_xml)


def write_analytics_sentiment(data, document):
    """
    Writes out tables for per-period, per-speaker sentiment from the analytics mode, as well as
    the overall sentiment for a speaker

    :param data: Transcribe results data
    :param document: Docx document to add the tables to
    """

    # Start with a new 2-column section
    document.add_section(WD_SECTION.CONTINUOUS)
    section_ptr = document.sections[-1]._sectPr
    cols = section_ptr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), '2')

    # Table 1 - Period sentiment per speaker
    write_custom_text_header(document, "Call Sentiment per Quarter of the call")
    table = document.add_table(rows=1, cols=5)
    table.style = document.styles[TABLE_STYLE_STANDARD]
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Speaker"
    hdr_cells[1].text = "Q1"
    hdr_cells[2].text = "Q2"
    hdr_cells[3].text = "Q3"
    hdr_cells[4].text = "Q4"
    for col in range(1, 5):
        hdr_cells[col].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Work through our sentiment period data
    period_sentiment = data["ConversationCharacteristics"]["Sentiment"]["SentimentByPeriod"]["QUARTER"]
    for caller in period_sentiment:
        # First column is the speaker
        row_cells = table.add_row().cells
        row_cells[0].text = caller.title()
        col_offset = 1
        # Further columns on that row hold the value for one period on the call
        for period in period_sentiment[caller]:
            row_cells[col_offset].text = str(period["Score"])
            row_cells[col_offset].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell_colour = get_text_colour_analytics_sentiment(period["Score"])
            set_table_cell_background_colour(row_cells[col_offset], cell_colour)
            col_offset += 1

    # Put in a short table footer, then move to the next column
    document.add_paragraph()  # Spacing
    write_small_header_text(document, "SENTIMENT: Range from +5 (Positive) to -5 (Negative)", 0.9)

    # Table 2 - Overall speaker sentiment
    write_custom_text_header(document, "Overall Speaker Sentiment")
    table = document.add_table(rows=1, cols=2)
    table.style = document.styles[TABLE_STYLE_STANDARD]
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Speaker"
    hdr_cells[1].text = "Sentiment"
    hdr_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    speaker_sentiment = data["ConversationCharacteristics"]["Sentiment"]["OverallSentiment"]
    for caller in speaker_sentiment:
        row_cells = table.add_row().cells
        row_cells[0].text = caller.title()
        row_cells[1].text = str(speaker_sentiment[caller])
        row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell_colour = get_text_colour_analytics_sentiment(speaker_sentiment[caller])
        set_table_cell_background_colour(row_cells[1], cell_colour)

    # Keep the columns narrow for the 2nd table
    widths = (Cm(2.2), Cm(1.5))
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width
    document.add_paragraph()  # Spacing


def create_turn_by_turn_segments(data, cli_args):
    """
    This creates a list of per-turn speech segments based upon the transcript data.  It has to work in three
    slightly different ways, as each operational mode from Transcribe outputs slightly different JSON structures.
    These modes are (a) Speaker-separated audio, (b) Channel-separated audio, and (c) Call Analytics audio

    :param data: JSON result data from Transcribe
    :param cli_args: CLI arguments used for this processing run
    :return: List of transcription speech segments
    :return: Flag to indicate the presence of call summary data
    """
    speechSegmentList = []
    summaries_detected = False

    # Decide on our operational mode - it's in the job-status or, if necessary, infer it from the data file
    # STANDARD => speaker separated, channel separated;  ANALYTICS => different format
    isAnalyticsMode = cli_args.analyticsMode
    if isAnalyticsMode:
        # We know if its analytics mode, as it's defined in the job-status and file
        isChannelMode = False
        isSpeakerMode = False
    else:
        # Channel/Speaker-mode only relevant if not using analytics
        isChannelMode = "channel_labels" in data["results"]
        isSpeakerMode = not isChannelMode

    lastSpeaker = ""
    lastEndTime = 0.0
    skipLeadingSpace = False
    confidenceList = []
    nextSpeechSegment = None

    # Process a Speaker-separated non-analytics file
    if isSpeakerMode:
        # A segment is a blob of pronunciation and punctuation by an individual speaker
        for segment in data["results"]["speaker_labels"]["segments"]:

            # If there is content in the segment then pick out the time and speaker
            if len(segment["items"]) > 0:
                # Pick out our next data
                nextStartTime = float(segment["start_time"])
                nextEndTime = float(segment["end_time"])
                nextSpeaker = str(segment["speaker_label"])

                # If we've changed speaker, or there's a gap, create a new row
                if (nextSpeaker != lastSpeaker) or ((nextStartTime - lastEndTime) >= START_NEW_SEGMENT_DELAY):
                    nextSpeechSegment = SpeechSegment()
                    speechSegmentList.append(nextSpeechSegment)
                    nextSpeechSegment.segmentStartTime = nextStartTime
                    nextSpeechSegment.segmentSpeaker = nextSpeaker
                    skipLeadingSpace = True
                    confidenceList = []
                    nextSpeechSegment.segmentConfidence = confidenceList
                nextSpeechSegment.segmentEndTime = nextEndTime

                # Note the speaker and end time of this segment for the next iteration
                lastSpeaker = nextSpeaker
                lastEndTime = nextEndTime

                # For each word in the segment...
                for word in segment["items"]:

                    # Get the word with the highest confidence
                    pronunciations = list(filter(lambda x: x["type"] == "pronunciation", data["results"]["items"]))
                    word_result = list(filter(lambda x: x["start_time"] == word["start_time"] and x["end_time"] == word["end_time"], pronunciations))
                    try:
                        result = sorted(word_result[-1]["alternatives"], key=lambda x: x["confidence"])[-1]
                        confidence = float(result["confidence"])
                    except:
                        result = word_result[-1]["alternatives"][0]
                        confidence = float(result["redactions"][0]["confidence"])

                    # Write the word, and a leading space if this isn't the start of the segment
                    if skipLeadingSpace:
                        skipLeadingSpace = False
                        wordToAdd = result["content"]
                    else:
                        wordToAdd = " " + result["content"]

                    # If the next item is punctuation, add it to the current word
                    try:
                        word_result_index = data["results"]["items"].index(word_result[0])
                        next_item = data["results"]["items"][word_result_index + 1]
                        if next_item["type"] == "punctuation":
                            wordToAdd += next_item["alternatives"][0]["content"]
                    except IndexError:
                        pass

                    nextSpeechSegment.segmentText += wordToAdd
                    confidenceList.append({"text": wordToAdd,
                                           "confidence": confidence,
                                           "start_time": float(word["start_time"]),
                                           "end_time": float(word["end_time"])})

    # Process a Channel-separated non-analytics file
    elif isChannelMode:

        # A channel contains all pronunciation and punctuation from a single speaker
        for channel in data["results"]["channel_labels"]["channels"]:

            # If there is content in the channel then start processing it
            if len(channel["items"]) > 0:

                # We have the same speaker all the way through this channel
                nextSpeaker = str(channel["channel_label"])
                for word in channel["items"]:
                    # Pick out our next data from a 'pronunciation'
                    if word["type"] == "pronunciation":
                        nextStartTime = float(word["start_time"])
                        nextEndTime = float(word["end_time"])

                        # If we've changed speaker, or we haven't and the
                        # pause is very small, then start a new text segment
                        if (nextSpeaker != lastSpeaker) or \
                                ((nextSpeaker == lastSpeaker) and ((nextStartTime - lastEndTime) > 0.1)):
                            nextSpeechSegment = SpeechSegment()
                            speechSegmentList.append(nextSpeechSegment)
                            nextSpeechSegment.segmentStartTime = nextStartTime
                            nextSpeechSegment.segmentSpeaker = nextSpeaker
                            skipLeadingSpace = True
                            confidenceList = []
                            nextSpeechSegment.segmentConfidence = confidenceList
                        nextSpeechSegment.segmentEndTime = nextEndTime

                        # Note the speaker and end time of this segment for the next iteration
                        lastSpeaker = nextSpeaker
                        lastEndTime = nextEndTime

                        # Get the word with the highest confidence
                        pronunciations = list(filter(lambda x: x["type"] == "pronunciation", channel["items"]))
                        word_result = list(filter(lambda x: x["start_time"] == word["start_time"] and x["end_time"] == word["end_time"], pronunciations))
                        try:
                            result = sorted(word_result[-1]["alternatives"], key=lambda x: x["confidence"])[-1]
                            confidence = float(result["confidence"])
                        except:
                            result = word_result[-1]["alternatives"][0]
                            confidence = float(result["redactions"][0]["confidence"])
                        # result = sorted(word_result[-1]["alternatives"], key=lambda x: x["confidence"])[-1]

                        # Write the word, and a leading space if this isn't the start of the segment
                        if (skipLeadingSpace):
                            skipLeadingSpace = False
                            wordToAdd = result["content"]
                        else:
                            wordToAdd = " " + result["content"]

                        # If the next item is punctuation, add it to the current word
                        try:
                            word_result_index = channel["items"].index(word_result[0])
                            next_item = channel["items"][word_result_index + 1]
                            if next_item["type"] == "punctuation":
                                wordToAdd += next_item["alternatives"][0]["content"]
                        except IndexError:
                            pass

                        # Finally, add the word and confidence to this segment's list
                        nextSpeechSegment.segmentText += wordToAdd
                        confidenceList.append({"text": wordToAdd,
                                               "confidence": confidence,
                                               "start_time": float(word["start_time"]),
                                               "end_time": float(word["end_time"])})

        # Sort the segments, as they are in channel-order and not speaker-order, then
        # merge together turns from the same speaker that are very close together
        speechSegmentList = sorted(speechSegmentList, key=lambda segment: segment.segmentStartTime)
        speechSegmentList = merge_speaker_segments(speechSegmentList)

    # Process a Call Analytics file
    elif isAnalyticsMode:

        # Lookup shortcuts
        interrupts = data["ConversationCharacteristics"]["Interruptions"]

        # Each turn has already been processed by Transcribe, so the outputs are in order
        for turn in data["Transcript"]:

            # Setup the next speaker block
            nextSpeechSegment = SpeechSegment()
            speechSegmentList.append(nextSpeechSegment)
            nextSpeechSegment.segmentStartTime = float(turn["BeginOffsetMillis"]) / 1000.0
            nextSpeechSegment.segmentEndTime = float(turn["EndOffsetMillis"]) / 1000.0
            nextSpeechSegment.segmentSpeaker = turn["ParticipantRole"].title()
            nextSpeechSegment.segmentText = turn["Content"]
            nextSpeechSegment.segmentLoudnessScores = turn["LoudnessScores"]
            confidenceList = []
            nextSpeechSegment.segmentConfidence = confidenceList
            skipLeadingSpace = True

            # Check if this block is within an interruption block for the speaker
            if turn["ParticipantRole"] in interrupts["InterruptionsByInterrupter"]:
                for entry in interrupts["InterruptionsByInterrupter"][turn["ParticipantRole"]]:
                    if turn["BeginOffsetMillis"] == entry["BeginOffsetMillis"]:
                        nextSpeechSegment.segmentInterruption = True

            # Record any issues detected
            if "IssuesDetected" in turn:
                summaries_detected = True
                for issue in turn["IssuesDetected"]:
                    # Grab the transcript offsets for the issue text
                    nextSpeechSegment.segmentIssuesDetected.append(issue["CharacterOffsets"])

            # Record any actions detected
            if "ActionItemsDetected" in turn:
                summaries_detected = True
                for action in turn["ActionItemsDetected"]:
                    # Grab the transcript offsets for the issue text
                    nextSpeechSegment.segmentActionItemsDetected.append(action["CharacterOffsets"])

            # Record any outcomes detected
            if "OutcomesDetected" in turn:
                summaries_detected = True
                for outcome in turn["OutcomesDetected"]:
                    # Grab the transcript offsets for the issue text
                    nextSpeechSegment.segmentOutcomesDetected.append(outcome["CharacterOffsets"])

            # Process each word in this turn
            for word in turn["Items"]:
                # Pick out our next data from a 'pronunciation'
                if word["Type"] == "pronunciation":
                    # Write the word, and a leading space if this isn't the start of the segment
                    if skipLeadingSpace:
                        skipLeadingSpace = False
                        wordToAdd = word["Content"]
                    else:
                        wordToAdd = " " + word["Content"]

                    # If the word is redacted then the word confidence is a bit more buried
                    if "Confidence" in word:
                        conf_score = float(word["Confidence"])
                    elif "Redaction" in word:
                        conf_score = float(word["Redaction"][0]["Confidence"])

                    # Add the word and confidence to this segment's list
                    confidenceList.append({"text": wordToAdd,
                                           "confidence": conf_score,
                                           "start_time": float(word["BeginOffsetMillis"]) / 1000.0,
                                           "end_time": float(word["BeginOffsetMillis"] / 1000.0)})
                else:
                    # Punctuation, needs to be added to the previous word
                    last_word = nextSpeechSegment.segmentConfidence[-1]
                    last_word["text"] = last_word["text"] + word["Content"]

            # Tag on the sentiment - analytics has no per-turn numbers
            turn_sentiment = turn["Sentiment"]
            if turn_sentiment == "POSITIVE":
                nextSpeechSegment.segmentIsPositive = True
                nextSpeechSegment.segmentPositive = 1.0
                nextSpeechSegment.segmentSentimentScore = 1.0
            elif turn_sentiment == "NEGATIVE":
                nextSpeechSegment.segmentIsNegative = True
                nextSpeechSegment.segmentNegative = 1.0
                nextSpeechSegment.segmentSentimentScore = 1.0

    # Return our full turn-by-turn speaker segment list with sentiment,
    # along with a flag to indicate the presence of call summary data
    return speechSegmentList, summaries_detected


def load_transcribe_job_status(cli_args):
    """
    Loads in the job status for the job named in cli_args.inputJob.  This will try both the standard Transcribe API
    as well as the Analytics API, as the customer may not know which one their job relates to

    :param cli_args: CLI arguments used for this processing run
    :return: The job status structure (different between standard/analytics), and a 'job-completed' flag
    """
    transcribe_client = boto3.client("transcribe")

    try:
        # Extract the standard Transcribe job status
        job_status = transcribe_client.get_transcription_job(TranscriptionJobName=cli_args.inputJob)["TranscriptionJob"]
        cli_args.analyticsMode = False
        completed = job_status["TranscriptionJobStatus"]
    except:
        # That job doesn't exist, but it may have been an analytics job
        job_status = transcribe_client.get_call_analytics_job(CallAnalyticsJobName=cli_args.inputJob)["CallAnalyticsJob"]
        cli_args.analyticsMode = True
        completed = job_status["CallAnalyticsJobStatus"]

    return job_status, completed


def generate_document():
    """
    Entrypoint for the command-line interface.
    """
    # Parameter extraction
    cli_parser = argparse.ArgumentParser(prog='ts-to-word',
                                         description='Turn an Amazon Transcribe job output into an MS Word document')
    source_group = cli_parser.add_mutually_exclusive_group(required=True)
    source_group.add_argument('--inputFile', metavar='filename', type=str, help='File containing Transcribe JSON output')
    source_group.add_argument('--inputJob', metavar='job-id', type=str, help='Transcribe job identifier')
    cli_parser.add_argument('--outputFile', metavar='filename', type=str, help='Output file to hold MS Word document')
    cli_parser.add_argument('--sentiment', choices=['on', 'off'], default='off', help='Enables sentiment analysis on each conversational turn via Amazon Comprehend')
    cli_parser.add_argument('--confidence', choices=['on', 'off'], default='off', help='Displays information on word confidence scores throughout the transcript')
    cli_parser.add_argument('--keep', action='store_true', help='Keeps any downloaded job transcript JSON file')
    cli_args = cli_parser.parse_args()

    # If we're downloading a job transcript then validate that we have a job, then download it
    if cli_args.inputJob is not None:
        try:
            job_info, job_status = load_transcribe_job_status(cli_args)
        except:
            # Exception, most-likely due to the job not existing
            print("NOT FOUND: Requested job-id '{0}' does not exist.".format(cli_args.inputJob))
            exit(-1)

        # If the job hasn't completed then there is no transcript available
        if job_status == "FAILED":
            print("{0}: Requested job-id '{1}' has failed to complete".format(job_status, cli_args.inputJob))
            exit(-1)
        elif job_status != "COMPLETED":
            print("{0}: Requested job-id '{1}' has not yet completed.".format(job_status, cli_args.inputJob))
            exit(-1)

        # The transcript is available from a signed URL - get the redacted if it exists, otherwise the non-redacted
        if "RedactedTranscriptFileUri" in job_info["Transcript"]:
            # Get the redacted transcript
            download_url = job_info["Transcript"]["RedactedTranscriptFileUri"]
        else:
            # Gen the non-redacted transcript
            download_url = job_info["Transcript"]["TranscriptFileUri"]
        cli_args.inputFile = cli_args.inputJob + "-asrOutput.json"

        # Try and download the JSON - this will fail if the job delivered it to
        # an S3 bucket, as in that case the service no longer has the results
        try:
            urllib.request.urlretrieve(download_url, cli_args.inputFile)
        except:
            print("UNAVAILABLE: Transcript for job-id '{0}' is not available for download.".format(cli_args.inputJob))
            exit(-1)

        # Set our output filename if one wasn't supplied
        if cli_args.outputFile is None:
            cli_args.outputFile = cli_args.inputJob + ".docx"

    # Load in the JSON file for processing
    json_filepath = Path(cli_args.inputFile)
    if json_filepath.is_file():
        json_data = json.load(open(json_filepath.absolute(), "r", encoding="utf-8"))
    else:
        print("FAIL: Specified JSON file '{0}' does not exists.".format(cli_args.inputFile))
        exit(-1)

    # If this is a file-input run then try and load the job status (which may no longer exist)
    if cli_args.inputJob is None:
        try:
            # Ensure we don't delete our JSON later, reset our output file to match the job-name if it's currently blank
            cli_args.keep = True
            if cli_args.outputFile is None:
                if "results" in json_data:
                    cli_args.outputFile = json_data["jobName"] + ".docx"
                    cli_args.inputJob = json_data["jobName"]
                else:
                    cli_args.outputFile = json_data["JobName"] + ".docx"
                    cli_args.inputJob = json_data["JobName"]
            job_info, job_status = load_transcribe_job_status(cli_args)
        except:
            # No job status - need to quickly work out what mode we're in,
            # as standard job results look different from analytical ones
            cli_args.inputJob = None
            cli_args.outputFile = cli_args.inputFile + ".docx"
            cli_args.analyticsMode = "results" not in json_data
            job_info = None

    # Disable Comprehend's sentiment if we're in Analytics mode
    if cli_args.analyticsMode:
        cli_args.sentiment = 'off'

    # Generate the core transcript
    start = perf_counter()
    speech_segments, summaries_detected = create_turn_by_turn_segments(json_data, cli_args)

    # Inject Comprehend-based sentiments into the segment list if required
    if cli_args.sentiment == 'on':
        # Work out the mapped language code, as Transcribe supports more languages than Comprehend.  Just
        # see if the Transcribe language code starts with any of those that Comprehend supports and use that
        sentiment_lang_code = None
        for comprehend_code in SENTIMENT_LANGUAGES:
            if job_info["LanguageCode"].startswith(comprehend_code):
                sentiment_lang_code = comprehend_code
                break

        # If we have no match then we cannot perform sentiment analysis
        if sentiment_lang_code is not None:
            generate_sentiment(speech_segments, sentiment_lang_code)
        else:
            cli_args.sentiment = 'off'

    # Write out our file and the performance statistics
    write(cli_args, speech_segments, job_info, summaries_detected)
    finish = perf_counter()
    duration = round(finish - start, 2)
    print(f"> Transcript {cli_args.outputFile} writen in {duration} seconds.")

    # Finally, remove any temporary downloaded JSON results file
    if (cli_args.inputJob is not None) and (not cli_args.keep):
        os.remove(cli_args.inputFile)

# Main entrypoint
def handler(event, context):
    print("Event:", event)
    print("Context:", context)
    generate_document()